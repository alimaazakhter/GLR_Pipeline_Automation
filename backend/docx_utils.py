from docx import Document
import re

# Fallback mapping if standard keys are used but template needs something else
TEMPLATE_TO_LLM_KEY = {
    'DATE_LOSS': "Date Taken",
    'INSURED_NAME': "Insured's Name",
    'INSURED_H_STREET': "Address of Loss",
    'CARRIER_NAME': "Carrier Name",
    'POLICY_NO': "Policy #",
    'SERVICE_PROVIDER': "Service Provider",
    'SERVICE_PROVIDER_ADDRESS': "Service Provider Address",
    'SERVICE_PROVIDER_PHONE': "Service Provider Phone",
}

def extract_placeholders_from_docx(docx_file):
    """
    Reads a DOCX file and extracts all placeholders of the formats:
    - {{KEY}}
    - [KEY]
    Returns a sorted list of unique placeholder names (uppercased and stripped).
    """
    try:
        doc = Document(docx_file)
        placeholders = set()
        # Patterns to find {{KEY}} or [KEY]
        patterns = [re.compile(r'\{\{(.*?)\}\}'), re.compile(r'\[(.*?)\]')]
        
        def find_in_text(text):
            if not text:
                return
            for pattern in patterns:
                for match in pattern.findall(text):
                    cleaned = match.strip()
                    # Filter out empty or obviously invalid matches
                    if cleaned and not any(c in cleaned for c in '\n\r\t'):
                        placeholders.add(cleaned)
                        
        # Read from paragraphs
        for para in doc.paragraphs:
            find_in_text(para.text)
            
        # Read from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        find_in_text(para.text)
                        
        return sorted(list(placeholders))
    except Exception as e:
        print(f"Error extracting placeholders from DOCX: {e}")
        return []

def fill_docx_template(template_file, key_value_pairs, output_path):
    """
    Fills a DOCX template with key-value pairs and saves the result to output_path.
    Resolves the 'split runs' bug by replacing text at the paragraph level and merging runs.
    """
    try:
        doc = Document(template_file)
        
        # Build case-insensitive key maps for easy lookup
        key_map = {k.lower(): v for k, v in key_value_pairs.items()}
        
        patterns = [re.compile(r'\{\{(.*?)\}\}'), re.compile(r'\[(.*?)\]')]
        warnings = set()
        
        def replace_placeholders(text):
            for pattern in patterns:
                matches = pattern.findall(text)
                for match in matches:
                    key = match.strip()
                    # Try direct match (case-insensitive) first
                    value = key_map.get(key.lower())
                    
                    # Fallback to map if needed
                    if value is None:
                        mapped_key = TEMPLATE_TO_LLM_KEY.get(key.upper())
                        if mapped_key:
                            value = key_map.get(mapped_key.lower())
                            
                    if value is not None:
                        # Perform the replacement
                        text = text.replace(f"{{{{{key}}}}}", str(value))
                        text = text.replace(f"[{key}]", str(value))
                    else:
                        warnings.add(key)
            return text

        def process_paragraph(para):
            full_text = "".join(run.text for run in para.runs)
            # Quick check if paragraph contains any brackets/curly braces before processing
            if not ('[' in full_text or '{{' in full_text):
                return
                
            new_text = replace_placeholders(full_text)
            if new_text != full_text:
                if para.runs:
                    # Update first run and empty subsequent runs to preserve paragraph level replacement 
                    # without breaking the document formatting or deleting the runs.
                    para.runs[0].text = new_text
                    for run in para.runs[1:]:
                        run.text = ""
                else:
                    para.text = new_text

        # Replace in main body paragraphs
        for para in doc.paragraphs:
            process_paragraph(para)
            
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        process_paragraph(para)
                        
        doc.save(output_path)
        if warnings:
            print(f"Warning: The following placeholders were not found in the extracted data: {sorted(warnings)}")
        return True
    except Exception as e:
        print(f"Error filling DOCX template: {e}")
        return False
